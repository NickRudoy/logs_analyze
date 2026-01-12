"""
AI-анализатор логов с использованием GigaChat API
Анализирует результаты из Excel и генерирует рекомендации

Установка зависимостей:
pip install gigachat pandas openpyxl python-docx

Использование:
python ai_analyzer.py report_20250112_143022.xlsx --auth-key YOUR_AUTH_KEY
"""

import argparse
import sys
from pathlib import Path
import pandas as pd
import json
from datetime import datetime
from typing import Dict, List, Any
import uuid

try:
    from gigachat import GigaChat
    HAS_GIGACHAT = True
except ImportError:
    HAS_GIGACHAT = False
    print("Ошибка: библиотека gigachat не установлена")
    print("Установите: pip install gigachat")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Предупреждение: python-docx не установлен. Генерация Word отчетов отключена.")
    print("Установите: pip install python-docx")


class ReportAnalyzer:
    """Анализатор Excel отчетов"""
    
    def __init__(self, report_path: Path):
        self.report_path = report_path
        self.data = {}
        self.load_report()
    
    def load_report(self):
        """Загружает данные из Excel"""
        print(f"Загрузка отчета: {self.report_path}")
        
        try:
            # Читаем все листы
            excel_file = pd.ExcelFile(self.report_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                self.data[sheet_name] = df
                print(f"  Загружен лист: {sheet_name} ({len(df)} строк)")
        
        except Exception as e:
            print(f"Ошибка при загрузке отчета: {e}")
            sys.exit(1)
    
    def get_summary(self) -> Dict[str, Any]:
        """Извлекает сводную информацию"""
        summary = {}
        
        if 'Сводка' in self.data:
            df = self.data['Сводка']
            for _, row in df.iterrows():
                summary[row['Метрика']] = row['Значение']
        
        return summary
    
    def get_suspicious_ips(self, limit: int = 20) -> List[Dict]:
        """Извлекает топ подозрительных IP"""
        if 'Подозрительные IP' not in self.data:
            return []
        
        df = self.data['Подозрительные IP']
        return df.head(limit).to_dict('records')
    
    def get_country_stats(self) -> List[Dict]:
        """Извлекает статистику по странам"""
        if 'Статистика по странам' not in self.data:
            return []
        
        df = self.data['Статистика по странам']
        return df.to_dict('records')
    
    def get_load_anomalies(self) -> List[Dict]:
        """Извлекает аномалии нагрузки"""
        if 'Аномалии нагрузки' not in self.data:
            return []
        
        df = self.data['Аномалии нагрузки']
        return df.to_dict('records')
    
    def get_datacenter_ips(self) -> List[Dict]:
        """Извлекает IP из датацентров"""
        if 'IP из датацентров' not in self.data:
            return []
        
        df = self.data['IP из датацентров']
        return df.to_dict('records')
    
    def prepare_analysis_context(self) -> Dict[str, Any]:
        """Подготавливает контекст для AI анализа"""
        return {
            'summary': self.get_summary(),
            'suspicious_ips': self.get_suspicious_ips(10),
            'country_stats': self.get_country_stats(),
            'load_anomalies': self.get_load_anomalies(),
            'datacenter_ips': self.get_datacenter_ips()
        }


class GigaChatAnalyzer:
    """Анализатор с использованием GigaChat API"""
    
    def __init__(self, auth_key: str, model: str = "GigaChat", verify_ssl: bool = False):
        """
        Инициализация клиента GigaChat
        
        Args:
            auth_key: Ключ авторизации (Authorization Key) из личного кабинета
            model: Модель GigaChat (GigaChat, GigaChat-Pro, GigaChat-Plus)
            verify_ssl: Проверка SSL сертификатов (по умолчанию False для физлиц)
        """
        self.auth_key = auth_key
        self.model = model
        self.verify_ssl = verify_ssl
        self.client = None
        self.initialize_client()
    
    def initialize_client(self):
        """Инициализирует клиент GigaChat"""
        try:
            self.client = GigaChat(
                credentials=self.auth_key,
                model=self.model,
                verify_ssl_certs=self.verify_ssl,
                scope="GIGACHAT_API_PERS"  # Для физических лиц
            )
            print(f"GigaChat инициализирован (модель: {self.model})")
            
            # Проверяем доступность API
            try:
                models = self.client.get_models()
                # Безопасное извлечение списка моделей
                if hasattr(models, 'data') and models.data:
                    model_list = []
                    for m in models.data:
                        # Пробуем разные варианты доступа к ID модели
                        if hasattr(m, 'id_'):
                            model_list.append(m.id_)
                        elif hasattr(m, 'id'):
                            model_list.append(m.id)
                        elif hasattr(m, 'model'):
                            model_list.append(m.model)
                        elif isinstance(m, dict):
                            model_list.append(m.get('id', m.get('model', str(m))))
                        else:
                            # Если это объект Pydantic, пробуем получить id_ через getattr
                            model_id = getattr(m, 'id_', None) or getattr(m, 'id', None)
                            if model_id:
                                model_list.append(model_id)
                            else:
                                model_list.append(str(m))
                    if model_list:
                        # Фильтруем только модели для чата (исключаем Embeddings)
                        chat_models = [m for m in model_list if 'Embeddings' not in str(m) and 'embedding' not in str(m).lower()]
                        if chat_models:
                            print(f"Доступные модели для анализа: {', '.join(chat_models[:5])}{'...' if len(chat_models) > 5 else ''}")
            except Exception as e:
                print(f"Предупреждение: не удалось получить список моделей: {e}")
                
        except Exception as e:
            print(f"Ошибка инициализации GigaChat: {e}")
            print("\nПроверьте:")
            print("1. Правильность ключа авторизации (Authorization Key)")
            print("2. Наличие купленных токенов в личном кабинете")
            print("3. Подключение к интернету")
            sys.exit(1)
    
    def analyze_security_threats(self, context: Dict[str, Any]) -> str:
        """Анализирует угрозы безопасности"""
        print("\nАнализ угроз безопасности...")
        
        prompt = f"""Ты - эксперт по кибербезопасности и анализу веб-трафика.

Проанализируй данные о подозрительном трафике на веб-сайте:

СВОДНАЯ ИНФОРМАЦИЯ:
{json.dumps(context['summary'], ensure_ascii=False, indent=2)}

ТОП ПОДОЗРИТЕЛЬНЫХ IP:
{json.dumps(context['suspicious_ips'][:5], ensure_ascii=False, indent=2)}

СТАТИСТИКА ПО СТРАНАМ:
{json.dumps(context['country_stats'][:5], ensure_ascii=False, indent=2)}

IP ИЗ ДАТАЦЕНТРОВ:
{json.dumps(context['datacenter_ips'][:5], ensure_ascii=False, indent=2)}

Проведи анализ и предоставь:
1. Оценку уровня угрозы (низкий/средний/высокий/критический)
2. Идентификацию типов атак (DDoS, сканирование, парсинг, боты и т.д.)
3. Конкретные IP или группы IP, требующие немедленного внимания
4. Признаки координированных атак или ботнетов
5. Оценку потенциального ущерба для сайта

Ответ должен быть структурированным и конкретным."""

        return self._send_request(prompt)
    
    def generate_recommendations(self, context: Dict[str, Any], threat_analysis: str) -> str:
        """Генерирует рекомендации по защите"""
        print("Генерация рекомендаций...")
        
        prompt = f"""На основе анализа угроз:

{threat_analysis}

И данных об аномалиях нагрузки:
{json.dumps(context['load_anomalies'][:3], ensure_ascii=False, indent=2) if context['load_anomalies'] else 'Нет данных'}

Предоставь конкретные, практические рекомендации:

1. НЕМЕДЛЕННЫЕ ДЕЙСТВИЯ (в течение часа):
   - Какие IP блокировать прямо сейчас
   - Какие правила файрвола добавить
   - Какие параметры сервера изменить

2. КРАТКОСРОЧНЫЕ МЕРЫ (1-7 дней):
   - Настройка систем защиты (WAF, rate limiting)
   - Мониторинг и алертинг
   - Обновление конфигураций

3. ДОЛГОСРОЧНАЯ СТРАТЕГИЯ:
   - Архитектурные улучшения
   - Внедрение защитных сервисов (Cloudflare, Qrator)
   - Процессы и процедуры реагирования

4. КОНКРЕТНЫЕ КОМАНДЫ И ПРАВИЛА:
   - Примеры команд iptables/nginx/apache
   - Правила блокировки по странам/ASN
   - Конфигурации rate limiting

Рекомендации должны быть максимально конкретными и применимыми."""

        return self._send_request(prompt)
    
    def analyze_business_impact(self, context: Dict[str, Any]) -> str:
        """Анализирует влияние на бизнес"""
        print("Анализ влияния на бизнес...")
        
        summary = context['summary']
        bounce_rate = summary.get('Процент отказов (%)', 'N/A')
        
        prompt = f"""Проанализируй влияние подозрительного трафика на бизнес-метрики:

ТЕКУЩАЯ СИТУАЦИЯ:
- Процент отказов: {bounce_rate}
- Всего записей в логе: {summary.get('Всего записей в логе', 'N/A')}
- Отказов: {summary.get('Отказов', 'N/A')}

ПОДОЗРИТЕЛЬНАЯ АКТИВНОСТЬ:
- Подозрительных IP: {len(context['suspicious_ips'])}
- IP из датацентров: {len(context['datacenter_ips'])}
- Аномалий нагрузки: {len(context['load_anomalies'])}

Оцени:
1. Влияние на пользовательский опыт и конверсию
2. Потери в SEO (как поисковики видят высокий bounce rate)
3. Нагрузку на инфраструктуру и связанные расходы
4. Риски для репутации бренда
5. Потенциальные финансовые потери

Предоставь количественные оценки там, где возможно, и обоснованные предположения."""

        return self._send_request(prompt)
    
    def create_executive_summary(self, threat_analysis: str, recommendations: str, business_impact: str) -> str:
        """Создает краткую сводку для руководства"""
        print("Создание executive summary...")
        
        prompt = f"""Создай краткую сводку (executive summary) для руководства компании на основе:

АНАЛИЗ УГРОЗ:
{threat_analysis}

РЕКОМЕНДАЦИИ:
{recommendations}

ВЛИЯНИЕ НА БИЗНЕС:
{business_impact}

Executive summary должен содержать:
1. Суть проблемы (2-3 предложения)
2. Уровень критичности
3. Ключевые цифры и факты
4. 3-5 главных действий, которые нужно предпринять
5. Ожидаемый результат от внедрения рекомендаций

Язык должен быть понятен нетехническим руководителям. Фокус на бизнес-ценность и риски."""

        return self._send_request(prompt)
    
    def _send_request(self, prompt: str, max_retries: int = 3) -> str:
        """Отправляет запрос в GigaChat с повторными попытками"""
        for attempt in range(max_retries):
            try:
                # Используем метод chat из SDK
                # Модель уже указана при инициализации клиента GigaChat
                response = self.client.chat(prompt)
                
                # Извлекаем текст ответа
                if hasattr(response, 'choices') and len(response.choices) > 0:
                    content = response.choices[0].message.content
                    
                    # Выводим информацию об использовании токенов
                    if hasattr(response, 'usage'):
                        usage = response.usage
                        print(f"  Токенов использовано: {usage.total_tokens} "
                              f"(запрос: {usage.prompt_tokens}, ответ: {usage.completion_tokens})")
                    
                    return content
                else:
                    return "Ошибка: пустой ответ от GigaChat"
            
            except Exception as e:
                error_msg = str(e)
                print(f"  Попытка {attempt + 1}/{max_retries} не удалась: {error_msg}")
                
                # Специфичные ошибки
                if "401" in error_msg or "Unauthorized" in error_msg:
                    print("  Ошибка авторизации. Проверьте ключ авторизации.")
                    if attempt == 0:
                        print("  Попробуйте получить новый ключ в личном кабинете.")
                    break
                elif "429" in error_msg or "rate limit" in error_msg.lower():
                    print("  Превышен лимит запросов. Ожидание...")
                    import time
                    time.sleep(5 * (attempt + 1))
                elif "insufficient" in error_msg.lower() or "balance" in error_msg.lower():
                    print("  Недостаточно токенов. Пополните баланс в личном кабинете.")
                    break
                else:
                    if attempt < max_retries - 1:
                        import time
                        time.sleep(2 ** attempt)
                    else:
                        return f"Ошибка при обращении к GigaChat после {max_retries} попыток: {error_msg}"
        
        return "Не удалось получить ответ от GigaChat"


class ReportGenerator:
    """Генератор итогового отчета"""
    
    def __init__(self, output_path: Path):
        self.output_path = output_path
    
    def generate_text_report(self, analysis_results: Dict[str, str], context: Dict[str, Any]):
        """Генерирует текстовый отчет"""
        report_file = self.output_path.with_suffix('.txt')
        
        print(f"\nГенерация текстового отчета: {report_file}")
        
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("AI-АНАЛИЗ БЕЗОПАСНОСТИ ВЕБ-ТРАФИКА\n")
            f.write("=" * 80 + "\n")
            f.write(f"Дата анализа: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Исходный отчет: {self.output_path.name}\n")
            f.write(f"Модель: GigaChat API\n")
            f.write("=" * 80 + "\n\n")
            
            f.write("EXECUTIVE SUMMARY\n")
            f.write("-" * 80 + "\n")
            f.write(analysis_results['executive_summary'] + "\n\n")
            
            f.write("АНАЛИЗ УГРОЗ БЕЗОПАСНОСТИ\n")
            f.write("-" * 80 + "\n")
            f.write(analysis_results['threat_analysis'] + "\n\n")
            
            f.write("ВЛИЯНИЕ НА БИЗНЕС\n")
            f.write("-" * 80 + "\n")
            f.write(analysis_results['business_impact'] + "\n\n")
            
            f.write("РЕКОМЕНДАЦИИ\n")
            f.write("-" * 80 + "\n")
            f.write(analysis_results['recommendations'] + "\n\n")
            
            f.write("=" * 80 + "\n")
            f.write("ПРИЛОЖЕНИЕ: СТАТИСТИКА\n")
            f.write("=" * 80 + "\n")
            f.write("\nСводная информация:\n")
            for key, value in context['summary'].items():
                f.write(f"  {key}: {value}\n")
        
        print(f"Текстовый отчет сохранен: {report_file}")
        return report_file
    
    def generate_word_report(self, analysis_results: Dict[str, str], context: Dict[str, Any]):
        """Генерирует Word отчет"""
        if not HAS_DOCX:
            print("Word отчет пропущен: python-docx не установлен")
            return None
        
        report_file = self.output_path.with_suffix('.docx')
        
        print(f"Генерация Word отчета: {report_file}")
        
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('AI-АНАЛИЗ БЕЗОПАСНОСТИ ВЕБ-ТРАФИКА', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Мета-информация
        doc.add_paragraph(f"Дата анализа: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Исходный отчет: {self.output_path.name}")
        doc.add_paragraph(f"Модель: GigaChat API")
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(analysis_results['executive_summary'])
        doc.add_page_break()
        
        # Анализ угроз
        doc.add_heading('Анализ угроз безопасности', 1)
        doc.add_paragraph(analysis_results['threat_analysis'])
        doc.add_page_break()
        
        # Влияние на бизнес
        doc.add_heading('Влияние на бизнес', 1)
        doc.add_paragraph(analysis_results['business_impact'])
        doc.add_page_break()
        
        # Рекомендации
        doc.add_heading('Рекомендации по защите', 1)
        doc.add_paragraph(analysis_results['recommendations'])
        doc.add_page_break()
        
        # Статистика
        doc.add_heading('Приложение: Статистика', 1)
        
        doc.add_heading('Сводная информация', 2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Метрика'
        hdr_cells[1].text = 'Значение'
        
        for key, value in context['summary'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
        
        doc.save(report_file)
        print(f"Word отчет сохранен: {report_file}")
        return report_file


def main():
    parser = argparse.ArgumentParser(
        description='AI-анализатор логов с использованием GigaChat API',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры использования:
  # Базовый анализ (для физических лиц)
  python ai_analyzer.py report.xlsx --auth-key YOUR_AUTH_KEY
  
  # С указанием модели GigaChat-Pro
  python ai_analyzer.py report.xlsx --auth-key YOUR_AUTH_KEY --model GigaChat-Pro
  
  # Только текстовый отчет
  python ai_analyzer.py report.xlsx --auth-key YOUR_AUTH_KEY --format txt
  
  # Для юридических лиц (с проверкой SSL)
  python ai_analyzer.py report.xlsx --auth-key YOUR_AUTH_KEY --scope GIGACHAT_API_CORP --verify-ssl

Получение ключа авторизации:
  1. Зайдите в личный кабинет Studio: https://developers.sber.ru/studio/
  2. Откройте проект GigaChat API
  3. Перейдите в раздел "Настройки API"
  4. Нажмите "Получить ключ" и скопируйте Authorization Key
        """
    )
    
    parser.add_argument('report_path', help='Путь к Excel отчету от analyze_direct_traffic.py')
    parser.add_argument('--auth-key', required=True, 
                       help='Ключ авторизации (Authorization Key) из личного кабинета GigaChat')
    parser.add_argument('--model', default='GigaChat', 
                       choices=['GigaChat', 'GigaChat-Pro', 'GigaChat-Plus', 'GigaChat-2', 
                               'GigaChat-2-Pro', 'GigaChat-2-Max', 'GigaChat-Max'],
                       help='Модель GigaChat. Модели 1-го поколения (GigaChat, GigaChat-Pro) автоматически перенаправляются на GigaChat-2, GigaChat-2-Pro. Рекомендуется GigaChat-2-Pro для лучшего качества анализа')
    parser.add_argument('--scope', default='GIGACHAT_API_PERS',
                       choices=['GIGACHAT_API_PERS', 'GIGACHAT_API_CORP'],
                       help='Scope для авторизации (PERS - физлица, CORP - юрлица)')
    parser.add_argument('--verify-ssl', action='store_true',
                       help='Включить проверку SSL сертификатов (для юрлиц)')
    parser.add_argument('--format', choices=['txt', 'docx', 'both'], default='both', 
                       help='Формат выходного отчета (по умолчанию: both)')
    parser.add_argument('--output', help='Путь для сохранения отчета (без расширения)')
    
    args = parser.parse_args()
    
    # Проверка файла отчета
    report_path = Path(args.report_path)
    if not report_path.exists():
        print(f"Ошибка: файл {report_path} не найден")
        sys.exit(1)
    
    # Определение имени выходного файла
    if args.output:
        output_path = Path(args.output)
    else:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = report_path.parent / f"ai_analysis_{timestamp}"
    
    print("=" * 80)
    print("AI-АНАЛИЗАТОР БЕЗОПАСНОСТИ ВЕБ-ТРАФИКА")
    print("Powered by GigaChat API")
    print("=" * 80)
    
    # Загрузка отчета
    analyzer = ReportAnalyzer(report_path)
    context = analyzer.prepare_analysis_context()
    
    print(f"\nЗагружено данных:")
    print(f"  Подозрительных IP: {len(context['suspicious_ips'])}")
    print(f"  Стран: {len(context['country_stats'])}")
    print(f"  Аномалий нагрузки: {len(context['load_anomalies'])}")
    print(f"  IP из датацентров: {len(context['datacenter_ips'])}")
    
    # Инициализация GigaChat
    print("\n" + "=" * 80)
    ai_analyzer = GigaChatAnalyzer(
        auth_key=args.auth_key,
        model=args.model,
        verify_ssl=args.verify_ssl
    )
    
    # Анализ
    print("=" * 80)
    print("ЭТАП 1: Анализ угроз безопасности")
    print("=" * 80)
    threat_analysis = ai_analyzer.analyze_security_threats(context)
    
    print("\n" + "=" * 80)
    print("ЭТАП 2: Анализ влияния на бизнес")
    print("=" * 80)
    business_impact = ai_analyzer.analyze_business_impact(context)
    
    print("\n" + "=" * 80)
    print("ЭТАП 3: Генерация рекомендаций")
    print("=" * 80)
    recommendations = ai_analyzer.generate_recommendations(context, threat_analysis)
    
    print("\n" + "=" * 80)
    print("ЭТАП 4: Создание executive summary")
    print("=" * 80)
    executive_summary = ai_analyzer.create_executive_summary(
        threat_analysis, recommendations, business_impact
    )
    
    # Сборка результатов
    analysis_results = {
        'threat_analysis': threat_analysis,
        'business_impact': business_impact,
        'recommendations': recommendations,
        'executive_summary': executive_summary
    }
    
    # Генерация отчетов
    print("\n" + "=" * 80)
    print("ГЕНЕРАЦИЯ ОТЧЕТОВ")
    print("=" * 80)
    
    generator = ReportGenerator(output_path)
    
    if args.format in ['txt', 'both']:
        generator.generate_text_report(analysis_results, context)
    
    if args.format in ['docx', 'both']:
        generator.generate_word_report(analysis_results, context)
    
    print("\n" + "=" * 80)
    print("АНАЛИЗ ЗАВЕРШЕН")
    print("=" * 80)
    print("\nВывод executive summary:\n")
    print(executive_summary)
    print("\n" + "=" * 80)


if __name__ == '__main__':
    main()
